from pathlib import Path

from langchain_core.documents import Document

from src.config import RAW_DOCS_DIR, SUPPORTED_EXTENSIONS
from src.document_processing.image_extractor import extract_images
from src.document_processing.ocr import run_ocr
from src.document_processing.table_extractor import extract_tables

try:
    from langchain_community.document_loaders import CSVLoader, Docx2txtLoader, PyPDFLoader, TextLoader
except ImportError:
    CSVLoader = None
    Docx2txtLoader = None
    PyPDFLoader = None
    TextLoader = None


def _chat_docs_dir(chat_id: str | None) -> Path:
    if chat_id:
        return RAW_DOCS_DIR / f"chat_{chat_id}"
    return RAW_DOCS_DIR


def _load_text(path: Path) -> list[Document]:
    if TextLoader is not None:
        for encoding in ("utf-8", "utf-8-sig", "cp1256", "latin-1"):
            try:
                return TextLoader(str(path), encoding=encoding).load()
            except UnicodeDecodeError:
                continue

    return [Document(page_content=path.read_text(encoding="utf-8", errors="replace"))]


def _load_docx(path: Path) -> list[Document]:
    if Docx2txtLoader is not None:
        try:
            return Docx2txtLoader(str(path)).load()
        except Exception:
            pass

    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise RuntimeError("Install python-docx or docx2txt to parse DOCX files.") from exc

    docx = DocxDocument(str(path))
    text = "\n".join(paragraph.text for paragraph in docx.paragraphs if paragraph.text.strip())
    return [Document(page_content=text, metadata={"source": str(path)})]


def load_file(path: str | Path) -> list[Document]:
    file_path = Path(path)
    suffix = file_path.suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        return []

    if suffix == ".pdf":
        if PyPDFLoader is None:
            raise RuntimeError("Install pypdf and langchain-community to parse PDF files.")
        loaded_docs = PyPDFLoader(str(file_path)).load()
    elif suffix == ".csv":
        if CSVLoader is not None:
            loaded_docs = CSVLoader(str(file_path), encoding="utf-8").load()
        else:
            loaded_docs = _load_text(file_path)
    elif suffix == ".docx":
        loaded_docs = _load_docx(file_path)
    else:
        loaded_docs = _load_text(file_path)

    for doc in loaded_docs:
        doc.metadata["source"] = str(file_path)
        doc.metadata["file_name"] = file_path.name
        doc.metadata["file_type"] = suffix
        doc.metadata["ocr"] = run_ocr(file_path).__dict__
        doc.metadata["image_extraction"] = extract_images(file_path).__dict__
        doc.metadata["table_extraction"] = extract_tables(file_path).__dict__

    return loaded_docs


def load_documents(
    chat_id: str | None = None,
    directory: str | Path | None = None,
    file_paths: list[str | Path] | None = None,
) -> list[Document]:
    if file_paths is None:
        docs_dir = Path(directory) if directory is not None else _chat_docs_dir(chat_id)
        if not docs_dir.exists():
            return []
        file_paths = [path for path in docs_dir.iterdir() if path.is_file()]

    documents: list[Document] = []
    for path in sorted(Path(p) for p in file_paths):
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        try:
            documents.extend(load_file(path))
        except Exception as exc:
            documents.append(
                Document(
                    page_content=f"Failed to load {path.name}: {exc}",
                    metadata={
                        "source": str(path),
                        "file_name": path.name,
                        "file_type": path.suffix.lower(),
                        "error": str(exc),
                    },
                )
            )

    return documents
